"""
Multi-format Document Processor for ClauseWise
Handles PDF, DOCX, and TXT files
"""

import os
import re
from typing import Optional, Dict, List
from pathlib import Path

class DocumentProcessor:
    """Process legal documents in multiple formats"""
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.txt']
    
    def __init__(self):
        self.text = ""
        self.filename = ""
        self.file_type = ""
        self.metadata = {}
    
    def process_file(self, file_path: str) -> Dict:
        """
        Process uploaded file and extract text
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_path = Path(file_path)
        self.filename = file_path.name
        self.file_type = file_path.suffix.lower()
        
        if self.file_type not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {self.file_type}. Supported: {self.SUPPORTED_FORMATS}")
        
        # Extract text based on file type
        if self.file_type == '.pdf':
            self.text = self._extract_from_pdf(file_path)
        elif self.file_type == '.docx':
            self.text = self._extract_from_docx(file_path)
        elif self.file_type == '.txt':
            self.text = self._extract_from_txt(file_path)
        
        # Clean and normalize text
        self.text = self._clean_text(self.text)
        
        # Extract basic metadata
        self.metadata = self._extract_metadata()
        
        return {
            'text': self.text,
            'filename': self.filename,
            'file_type': self.file_type,
            'metadata': self.metadata
        }
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import pdfplumber
            
            text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            
            return "\n\n".join(text)
        except Exception as e:
            # Fallback to PyPDF2
            try:
                import PyPDF2
                text = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text.append(page.extract_text())
                return "\n\n".join(text)
            except Exception as e2:
                raise Exception(f"Error extracting PDF: {e}, {e2}")
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text.append(row_text)
            
            return "\n\n".join(text)
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {e}")
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['latin-1', 'iso-8859-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except:
                    continue
            raise Exception("Error: Unable to decode text file with any known encoding")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers (common patterns)
        text = re.sub(r'Page\s+\d+\s+of\s+\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\bPage\s+\d+\b', '', text, flags=re.IGNORECASE)
        
        # Remove excessive line breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()
    
    def _extract_metadata(self) -> Dict:
        """Extract basic metadata from document"""
        words = len(self.text.split())
        chars = len(self.text)
        
        # Count potential clauses (heuristic: sentences ending with period/semicolon)
        sentences = len(re.findall(r'[.;]\s+', self.text))
        
        # Detect if it contains common legal keywords
        legal_keywords = [
            'whereas', 'hereby', 'hereunder', 'hereinafter',
            'notwithstanding', 'pursuant', 'covenant', 'indemnify',
            'liability', 'breach', 'termination', 'confidential'
        ]
        
        keyword_count = sum(1 for keyword in legal_keywords 
                          if keyword.lower() in self.text.lower())
        
        return {
            'word_count': words,
            'char_count': chars,
            'estimated_clauses': sentences,
            'legal_keyword_density': keyword_count / len(legal_keywords) * 100,
            'file_size': len(self.text.encode('utf-8'))
        }
    
    def chunk_text(self, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """
        Split text into overlapping chunks for processing
        
        Args:
            chunk_size: Maximum size of each chunk in characters
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not self.text:
            return []
        
        chunks = []
        start = 0
        text_length = len(self.text)
        
        while start < text_length:
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence ending
                sentence_end = self.text.rfind('.', start, end)
                if sentence_end > start:
                    end = sentence_end + 1
            
            chunks.append(self.text[start:end].strip())
            start = end - overlap
        
        return chunks
    
    def get_document_summary(self) -> str:
        """Get a brief summary of the document"""
        summary = f"""
        Document: {self.filename}
        Type: {self.file_type.upper()}
        Word Count: {self.metadata.get('word_count', 0):,}
        Estimated Clauses: {self.metadata.get('estimated_clauses', 0)}
        Legal Density: {self.metadata.get('legal_keyword_density', 0):.1f}%
        """
        return summary.strip()